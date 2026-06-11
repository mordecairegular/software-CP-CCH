#!/usr/bin/env python3
"""Build final DOCX materials for a Chinese software copyright case.

The tool deliberately separates verified case facts from placeholders. If a
field is still pending, it writes a visible confirmation marker instead of
leaking values from a sample template.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:  # pragma: no cover - optional for structural generation
    Image = None
    ImageDraw = None
    ImageFont = None


LINES_PER_PAGE = 50
FRONT_LINES = 1650
BACK_LINES = 1650


@dataclass
class GeneratedFile:
    label: str
    path: Path


def read_text(path: Path) -> str:
    if not path.exists():
        return ""
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk", "cp936"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def read_json_any(path: Path):
    if not path.exists():
        return None
    try:
        return json.loads(read_text(path))
    except Exception:
        return None


def count_cjk(text: str) -> int:
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def text_quality_issues(label: str, text: str, *, min_cjk: int = 30) -> list[str]:
    issues: list[str] = []
    replacement_count = text.count("\ufffd")
    qmark_runs = re.findall(r"\?{3,}", text)
    qmark_count = text.count("?")
    if replacement_count:
        issues.append(f"FAIL: {label} contains replacement characters, count={replacement_count}")
    if qmark_runs and qmark_count > 20:
        issues.append(f"FAIL: {label} contains suspicious question-mark mojibake, count={qmark_count}")
    if len(text) > 200 and min_cjk and count_cjk(text) < min_cjk:
        issues.append(f"FAIL: {label} has too few Chinese characters after generation; possible encoding loss")
    return issues


def safe_filename(text: str) -> str:
    text = re.sub(r'[<>:"/\\|?*\n\r\t]+', "_", text)
    return re.sub(r"\s+", " ", text).strip()[:110]


def clean_value(value: str) -> str:
    value = value.strip()
    value = re.sub(r"\s+", " ", value)
    return value.strip("：: ")


def parse_application_fields(path: Path) -> dict[str, str]:
    data: dict[str, str] = {}
    text = read_text(path)
    for line in text.splitlines():
        if not line.strip().startswith("➤"):
            continue
        line = line.strip().lstrip("➤").strip()
        if "：" in line:
            key, value = line.split("：", 1)
        elif ":" in line:
            key, value = line.split(":", 1)
        else:
            continue
        key = clean_value(key).replace(" / ", "/")
        data[key] = clean_value(value)
    return data


def pick_field(fields: dict[str, str], names: Iterable[str], default: str) -> str:
    for name in names:
        if name in fields and fields[name].strip():
            return fields[name].strip()
    return default


def pending_marker(label: str) -> str:
    return f"【待确认：{label}】"


def is_pending(value: str) -> bool:
    return (not value.strip()) or ("待" in value) or ("确认" in value)


def load_metadata(workdir: Path, args: argparse.Namespace) -> dict[str, str]:
    fields = parse_application_fields(workdir / "草稿" / "申请表信息.md")
    business_json: dict[str, str] = {}
    business_json_path = workdir / "草稿" / "业务理解.json"
    if business_json_path.exists():
        try:
            loaded = json.loads(read_text(business_json_path))
            if isinstance(loaded, dict):
                business_json = {str(key): str(value) for key, value in loaded.items() if not isinstance(value, (list, dict))}
        except json.JSONDecodeError:
            business_json = {}
    code_manifest = workdir / "草稿" / "代码提取清单.json"
    source_lines = ""
    if code_manifest.exists():
        try:
            manifest_data = json.loads(read_text(code_manifest))
            source_lines = str(sum(int(item.get("line_count") or 0) for item in manifest_data.get("files", [])))
            if source_lines == "0":
                source_lines = str(manifest_data.get("total_lines", ""))
        except json.JSONDecodeError:
            source_lines = ""

    software_name = args.software_name or pick_field(fields, ["软件全称"], pending_marker("软件全称"))
    short_name = args.short_name or pick_field(fields, ["软件简称"], "")
    version = args.version or pick_field(fields, ["版本号"], "V1.0")

    meta = {
        "software_name": software_name,
        "short_name": short_name or software_name,
        "version": version,
        "owner": args.owner or pick_field(fields, ["著作权人/申请人", "著作权人 / 申请人", "著作权人"], pending_marker("著作权人")),
        "developer": args.developer or pick_field(fields, ["开发者"], pending_marker("开发者")),
        "completion_date": args.completion_date or pick_field(fields, ["开发完成日期"], pending_marker("开发完成日期")),
        "first_publish": args.first_publish or pick_field(fields, ["首次发表日期"], "未发表（待确认）"),
        "rights_acquire": args.rights_acquire or pick_field(fields, ["权利取得方式"], "原始取得（待确认）"),
        "development_mode": args.development_mode or pick_field(fields, ["开发方式"], "独立开发/职务开发待确认"),
        "dev_hardware": pick_field(fields, ["开发硬件环境"], "普通个人计算机，8GB 及以上内存"),
        "runtime_hardware": pick_field(fields, ["运行硬件环境"], "普通个人计算机，本地运行环境"),
        "dev_os": pick_field(fields, ["开发操作系统"], "Windows"),
        "runtime_os": pick_field(fields, ["运行平台/操作系统"], "Windows"),
        "dev_tools": pick_field(fields, ["开发工具", "软件开发环境/开发工具"], "Python、文本编辑器、版本管理工具"),
        "support_env": pick_field(fields, ["运行支撑环境", "软件运行支撑环境/支持软件"], "Python 3.10 以上及项目依赖库"),
        "language": pick_field(fields, ["编程语言"], "Python"),
        "source_lines": args.source_lines or source_lines or pick_field(fields, ["源程序量"], pending_marker("源程序量")),
        "software_class": pick_field(fields, ["软件分类"], "应用软件"),
        "industry": args.industry or pick_field(fields, ["行业领域", "面向领域/行业"], business_json.get("industry") or "待确认"),
        "classification_code": args.classification_code or "30200/6510（待确认）",
        "owner_category": args.owner_category or "待确认",
        "certificate_type": args.certificate_type or "待确认",
        "certificate_no": args.certificate_no or "【待确认：证件号码】",
        "nationality": args.nationality or "中国",
        "province_city": args.province_city or "【待确认：省份/城市】",
        "contact_name": args.contact_name or pending_marker("联系人"),
        "contact_phone": args.contact_phone or pending_marker("电话"),
        "contact_mobile": args.contact_mobile or pending_marker("手机"),
        "contact_email": args.contact_email or pending_marker("Email"),
        "contact_address": args.contact_address or pending_marker("详细地址"),
        "contact_postcode": args.contact_postcode or pending_marker("邮编"),
    }
    return meta


def extract_business_paragraphs(workdir: Path) -> dict[str, str]:
    text = read_text(workdir / "草稿" / "业务理解.md")
    manual = read_text(workdir / "草稿" / "操作手册.md")
    joined = "\n".join([text, manual])

    def section(title: str) -> str:
        pattern = re.compile(rf"^#+\s*{re.escape(title)}\s*$([\s\S]*?)(?=^#+\s|\Z)", re.MULTILINE)
        match = pattern.search(joined)
        return clean_markdown(match.group(1)) if match else ""

    return {
        "positioning": section("软件定位") or section("软件概述"),
        "core_problem": section("核心问题和使用场景"),
        "functions": section("主要功能模块") or section("主要功能"),
        "flow": section("典型操作流程"),
        "technical": section("技术特点的申请表口径") or section("技术特点"),
        "environment": section("运行环境"),
        "notes": section("使用注意事项"),
    }


def clean_markdown(text: str) -> str:
    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^\s*[-*]\s+", "・", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_mermaid_blocks(text: str) -> list[str]:
    return [match.group(1).strip() for match in re.finditer(r"```mermaid\s*([\s\S]*?)```", text, flags=re.IGNORECASE)]


def compose_main_function_text(meta: dict[str, str], business: dict[str, str]) -> str:
    base = business.get("positioning") or ""
    functions = business.get("functions") or ""
    flow = business.get("flow") or ""
    problem = business.get("core_problem") or ""
    industry = strip_terminal_punct(meta["industry"])
    pieces = [f"{meta['software_name']}是一款面向{industry}的应用软件。"]
    if base:
        pieces.append(base)
    if problem:
        pieces.append(problem)
    if functions:
        pieces.append(f"主要功能包括：{functions}")
    else:
        pieces.append("主要功能包括数据导入、参数配置、业务处理、结果校验、日志反馈和成果导出。")
    if flow:
        pieces.append(f"典型使用流程为：{flow}")
    pieces.append(
        "软件围绕真实业务对象、输入数据、处理规则和输出成果组织功能，"
        "用于把重复的数据处理、校验、生成和归档工作转化为可执行、可复核、可留痕的软件流程。"
    )
    text = " ".join(pieces)
    return re.sub(r"\s+", " ", text).strip()


def compose_technical_text(meta: dict[str, str], business: dict[str, str]) -> str:
    technical = business.get("technical") or ""
    environment = business.get("environment") or ""
    pieces = []
    if technical:
        pieces.append(technical)
    pieces.append(
        "软件按数据导入、参数配置、业务处理、结果生成、日志反馈和文件导出等环节组织程序模块，"
        "通过真实项目代码中的界面、数据解析、业务处理和成果输出逻辑完成主要功能。"
    )
    if environment:
        pieces.append(f"运行环境口径：{environment}")
    pieces.append(f"编程语言和支撑环境以申请表记录为准：{meta['language']}；{meta['support_env']}。")
    text = " ".join(pieces)
    return re.sub(r"\s+", " ", text).strip()


def strip_terminal_punct(text: str) -> str:
    return re.sub(r"[。；;，,\s]+$", "", text.strip())


def iter_unique_cells(row):
    seen = set()
    for cell in row.cells:
        key = cell._tc
        if key in seen:
            continue
        seen.add(key)
        yield cell


def unique_cells(row):
    return list(iter_unique_cells(row))


def set_cell_text(cell, text: str, *, size: int = 9, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_row_cell(table, row_idx: int, unique_idx: int, value: str, expected: str, report: list[str]) -> None:
    cells = unique_cells(table.rows[row_idx])
    row_text = " ".join(c.text for c in cells)
    if expected not in row_text:
        report.append(f"WARN: row {row_idx} expected label {expected!r}, actual={row_text[:80]!r}")
    if unique_idx >= len(cells):
        report.append(f"FAIL: row {row_idx} has no unique cell {unique_idx} for {expected}")
        return
    set_cell_text(cells[unique_idx], value)


def find_application_template(template_dir: Path | None) -> Path | None:
    if template_dir is None or not template_dir.exists():
        return None
    candidates = list(template_dir.glob("*.docx"))
    preferred = [
        p
        for p in candidates
        if ("申请表" in p.name or "登记申请" in p.name) and not p.name.startswith("~$")
    ]
    if preferred:
        return sorted(preferred)[0]
    return None


def reference_manual_image_count(reference_word_dir: Path | None) -> int:
    if reference_word_dir is None or not reference_word_dir.exists():
        return 0
    candidates = [p for p in sorted(reference_word_dir.glob("03 *.docx")) if not p.name.startswith("~$")]
    if not candidates:
        candidates = [p for p in sorted(reference_word_dir.glob("*.docx")) if "说明书" in p.name and not p.name.startswith("~$")]
    if not candidates:
        return 0
    try:
        doc = Document(str(candidates[0]))
    except Exception:
        return 0
    return sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)


def build_application_docx(workdir: Path, out_dir: Path, meta: dict[str, str], business: dict[str, str], template_dir: Path) -> tuple[GeneratedFile, list[str]]:
    qa: list[str] = []
    template = find_application_template(template_dir)
    out_path = out_dir / f"01 计算机软件著作权登记申请表 — {safe_filename(meta['software_name'] + meta['version'])}.docx"
    if template:
        shutil.copyfile(template, out_path)
        doc = Document(str(out_path))
        if not doc.tables:
            qa.append("FAIL: application template has no table; fallback form required")
        else:
            table = doc.tables[0]
            fill_application_table(table, meta, business, qa)
    else:
        qa.append("WARN: no application template found; generated simple structured form")
        doc = create_simple_application_doc(meta, business)
    doc.save(str(out_path))
    qa.extend(verify_application_docx(out_path, meta))
    return GeneratedFile("申请表", out_path), qa


def fill_application_table(table, meta: dict[str, str], business: dict[str, str], qa: list[str]) -> None:
    main_functions = compose_main_function_text(meta, business)
    technical = compose_technical_text(meta, business)
    rows_needed = 30
    if len(table.rows) < rows_needed:
        qa.append(f"WARN: application table row count {len(table.rows)} < {rows_needed}; template may differ")

    set_row_cell(table, 0, 2, meta["software_name"], "软件名称", qa)
    set_row_cell(table, 0, 4, meta["version"], "版本号", qa)
    set_row_cell(table, 1, 2, meta["short_name"], "软件简称", qa)
    set_row_cell(table, 1, 4, meta["classification_code"], "分类号", qa)
    set_row_cell(table, 2, 2, "●原创 ○修改（含翻译软件、合成软件）", "软件作品说明", qa)
    set_row_cell(table, 3, 1, meta["completion_date"], "开发完成日期", qa)
    published = "○已发表 首次发表日期：______ 首次发表地点及其所在国：______  ●未发表（☑允许公众查询）"
    if "已发表" in meta["first_publish"] and "未发表" not in meta["first_publish"]:
        published = f"●已发表 首次发表日期：{meta['first_publish']} 首次发表地点及其所在国：【待确认】  ○未发表"
    set_row_cell(table, 4, 1, published, "发表状态", qa)
    dev_mode = "●独立开发 ○合作开发 ○委托开发 ○下达任务完成"
    if "职务" in meta["development_mode"]:
        dev_mode = "●独立开发（职务开发事实待确认） ○合作开发 ○委托开发 ○下达任务完成"
    set_row_cell(table, 5, 1, dev_mode, "开发方式", qa)
    set_row_cell(table, 7, 1, meta["owner"], "著作权人", qa)
    set_row_cell(table, 7, 2, meta["owner_category"], "著作权人", qa)
    set_row_cell(table, 7, 3, meta["certificate_type"], "著作权人", qa)
    set_row_cell(table, 7, 4, meta["certificate_no"], "著作权人", qa)
    set_row_cell(table, 7, 5, meta["nationality"], "著作权人", qa)
    set_row_cell(table, 7, 6, meta["province_city"], "著作权人", qa)
    set_row_cell(table, 8, 2, "●原始取得 ○继续取得（受让 承受 继承）", "权利取得方式", qa)
    set_row_cell(table, 9, 2, "●全部 ○部分", "权利范围", qa)
    set_row_cell(table, 10, 1, "☑应用软件 □嵌入式软件 □中间件 □操作系统", "软件分类", qa)
    set_row_cell(table, 11, 2, "●一般交存：提交源程序鉴别材料；提交一种软件文档（操作说明书）", "一般交存", qa)
    set_row_cell(table, 13, 2, truncate(meta["dev_hardware"], 50), "开发的硬件环境", qa)
    set_row_cell(table, 14, 2, truncate(meta["runtime_hardware"], 50), "运行的硬件环境", qa)
    set_row_cell(table, 15, 2, truncate(meta["dev_os"], 50), "操作系统", qa)
    set_row_cell(table, 16, 2, truncate(meta["dev_tools"], 50), "开发环境", qa)
    set_row_cell(table, 17, 2, truncate(meta["runtime_os"], 50), "运行平台", qa)
    set_row_cell(table, 18, 2, truncate(meta["support_env"], 50), "运行支撑环境", qa)
    set_row_cell(table, 19, 2, meta["language"], "编程语言", qa)
    set_row_cell(table, 19, 4, only_digits_or_text(meta["source_lines"]), "源程序量", qa)
    industry = strip_terminal_punct(meta["industry"])
    purpose = f"用于{industry}中的业务处理、数据校验、结果生成和资料归档。"
    set_row_cell(table, 20, 2, truncate(purpose, 50), "开发目的", qa)
    set_row_cell(table, 21, 2, truncate(industry, 50), "面向领域", qa)
    set_row_cell(table, 22, 2, main_functions, "主要功能", qa)
    set_row_cell(table, 23, 2, "○APP ○游戏软件 ○教育软件 ○金融软件 ○医疗软件 ○地理信息软件 ○云计算软件 ○信息安全软件 ○大数据软件 ○人工智能软件 ○物联网软件 ●其他", "技术类别", qa)
    set_row_cell(table, 24, 2, truncate(technical, 100), "技术特点", qa)
    set_row_cell(table, 25, 2, "●由著作权人申请 ○由代理人申请（待确认）", "申请方式", qa)
    set_row_cell(table, 26, 2, meta["contact_name"], "姓名或名称", qa)
    set_row_cell(table, 26, 4, meta["contact_phone"], "电话", qa)
    set_row_cell(table, 27, 2, meta["contact_address"], "详细地址", qa)
    set_row_cell(table, 27, 4, meta["contact_postcode"], "邮编", qa)
    set_row_cell(table, 28, 2, meta["contact_name"], "联系人", qa)
    set_row_cell(table, 28, 4, meta["contact_mobile"], "手机", qa)
    set_row_cell(table, 29, 2, meta["contact_email"], "Email", qa)
    set_row_cell(table, 29, 4, "", "传真", qa)


def truncate(text: str, limit: int) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text if len(text) <= limit else text[: limit - 1] + "…"


def only_digits_or_text(text: str) -> str:
    match = re.search(r"\d+", text)
    if match:
        return match.group(0)
    return text


def collect_document_text(doc: Document) -> str:
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in unique_cells(row)))
    for section in doc.sections:
        parts.extend(p.text for p in section.header.paragraphs)
        parts.extend(p.text for p in section.footer.paragraphs)
        for table in list(section.header.tables) + list(section.footer.tables):
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in unique_cells(row)))
    return "\n".join(parts)


def create_simple_application_doc(meta: dict[str, str], business: dict[str, str]) -> Document:
    doc = Document()
    setup_normal_style(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("计算机软件著作权登记申请表")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(18)
    run.bold = True
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("软件名称", meta["software_name"]),
        ("软件简称", meta["short_name"]),
        ("版本号", meta["version"]),
        ("著作权人", meta["owner"]),
        ("开发者", meta["developer"]),
        ("开发完成日期", meta["completion_date"]),
        ("首次发表日期", meta["first_publish"]),
        ("主要功能", compose_main_function_text(meta, business)),
        ("技术特点", compose_technical_text(meta, business)),
    ]:
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
    return doc


def verify_application_docx(path: Path, meta: dict[str, str]) -> list[str]:
    qa: list[str] = []
    doc = Document(str(path))
    all_text = collect_document_text(doc)
    qa.extend(text_quality_issues("application docx text", all_text, min_cjk=60))
    for label, value in [
        ("软件名称", meta["software_name"]),
        ("版本号", meta["version"]),
        ("著作权人", meta["owner"]),
    ]:
        if value and value not in all_text:
            qa.append(f"FAIL: application field not found after save: {label}={value}")
    leaked = [
        term.strip()
        for term in os.environ.get("SOFTWARE_CP_CCH_LEAK_TERMS", "").split("|")
        if term.strip()
    ]
    for item in leaked:
        if item in all_text and item not in {meta["software_name"], meta["short_name"], meta["contact_name"], meta["contact_phone"], meta["contact_email"]}:
            qa.append(f"WARN: sample value still present in application form: {item}")
    return qa or ["PASS: application key fields found after save"]


def setup_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.0)


def add_heading_cn(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.style = f"Heading {min(level, 3)}"
    run = para.add_run(text)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    run.bold = True


def set_header_footer(doc: Document, title: str) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = title
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_field(footer)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页")
    run2.font.size = Pt(9)


def load_code_manifest(workdir: Path) -> dict:
    manifest_path = workdir / "草稿" / "代码提取清单.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"missing code manifest: {manifest_path}")
    return json.loads(read_text(manifest_path))


def read_code_lines(project_root: Path, manifest: dict) -> list[tuple[str, int, str]]:
    lines: list[tuple[str, int, str]] = []
    for item in manifest.get("files", []):
        rel = item["path"]
        path = project_root / rel
        if not path.exists():
            continue
        file_lines = read_text(path).splitlines()
        start = int(item.get("start_line") or 1)
        end = item.get("end_line") or len(file_lines)
        for idx in range(start, int(end) + 1):
            text = file_lines[idx - 1] if 0 <= idx - 1 < len(file_lines) else ""
            lines.append((rel, idx, text))
    return lines


def select_code_lines(lines: list[tuple[str, int, str]]) -> tuple[str, list[tuple[str, int, str]]]:
    if len(lines) <= FRONT_LINES + BACK_LINES:
        return "all-source-under-threshold", lines
    return "front-back-general-deposit", lines[:FRONT_LINES] + lines[-BACK_LINES:]


def build_source_docx(workdir: Path, out_dir: Path, meta: dict[str, str]) -> tuple[GeneratedFile, list[str]]:
    qa: list[str] = []
    manifest = load_code_manifest(workdir)
    project_root = Path(manifest.get("project_root") or workdir)
    lines = read_code_lines(project_root, manifest)
    mode, selected = select_code_lines(lines)
    out_path = out_dir / f"02 软件源代码 — {safe_filename(meta['short_name'])} {safe_filename(meta['version'])}.docx"

    doc = Document()
    setup_normal_style(doc)
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.6)
    set_header_footer(doc, f"{meta['software_name']}{meta['version']} 源程序鉴别材料")
    add_cover(doc, meta, "软件源代码", f"生成方式：{'全量提交' if mode.startswith('all') else '前后连续段提交'}；真实源码行数：{len(lines)}；本文件代码行数：{len(selected)}")

    current_file = None
    code_line_count = 0
    page_no = 1
    add_heading_cn(doc, f"第 {page_no} 页", 1)
    add_code_meta_line(doc, f"软件名称：{meta['software_name']}    版本号：{meta['version']}")
    for rel, line_no, text in selected:
        if rel != current_file:
            current_file = rel
            add_code_meta_line(doc, f"// FILE: {rel} lines from {line_no}", bold=True)
        add_code_line(doc, text)
        code_line_count += 1
        if code_line_count % LINES_PER_PAGE == 0 and code_line_count < len(selected):
            doc.add_page_break()
            page_no += 1
            add_heading_cn(doc, f"第 {page_no} 页", 1)
            add_code_meta_line(doc, f"软件名称：{meta['software_name']}    版本号：{meta['version']}")
    doc.save(str(out_path))
    qa.append(f"PASS: source doc generated, mode={mode}, real_lines={len(lines)}, selected_lines={len(selected)}, target_pages={max(1, (len(selected)+LINES_PER_PAGE-1)//LINES_PER_PAGE)}")
    if len(lines) < 3000:
        qa.append("WARN: real source code is under 3000 lines; generated all-source document rather than padding or fabricating code")
    return GeneratedFile("源代码", out_path), qa


def add_cover(doc: Document, meta: dict[str, str], doc_type: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(f"{meta['software_name']}{meta['version']}")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(20)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(doc_type)
    run2.font.name = "黑体"
    run2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run2.font.size = Pt(22)
    run2.bold = True
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(subtitle)
    run3.font.size = Pt(10.5)
    doc.add_page_break()


def add_code_meta_line(doc: Document, text: str, *, bold: bool = False) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.bold = bold


def add_code_line(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text if text else " ")
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)


def find_screenshots(workdir: Path, explicit_dir: Path | None) -> list[Path]:
    manifest_paths: list[Path] = []
    if explicit_dir:
        manifest_paths.append(explicit_dir / "screenshot_manifest.json")
    manifest_paths.append(workdir / "截图" / "screenshot_manifest.json")
    for manifest_path in manifest_paths:
        manifest_files = screenshots_from_manifest(manifest_path)
        if manifest_files is not None:
            return manifest_files

    if explicit_dir:
        dirs = [explicit_dir]
    else:
        word_dir = workdir / "正式资料" / "word" / "screenshots"
        dirs = [word_dir] if list(word_dir.glob("*.png")) else [workdir / "截图"]
    seen = set()
    files: list[Path] = []
    for directory in dirs:
        if not directory.exists():
            continue
        for path in sorted(directory.glob("*.png")):
            if path.resolve() in seen:
                continue
            seen.add(path.resolve())
            files.append(path)
    preferred_terms = ["word_", "desktop", "batch", "history", "report", "mobile", "curve"]
    files.sort(key=lambda p: (min([p.name.find(t) if t in p.name else 999 for t in preferred_terms]), p.name))
    return files


def screenshots_from_manifest(manifest_path: Path) -> list[Path] | None:
    manifest = read_json_any(manifest_path)
    if manifest is None:
        return None
    if isinstance(manifest, dict):
        records = manifest.get("screenshots") or manifest.get("items") or []
    elif isinstance(manifest, list):
        records = manifest
    else:
        records = []
    selected: list[tuple[int, str, Path]] = []
    for index, item in enumerate(records):
        if not isinstance(item, dict):
            continue
        if item.get("accepted") is False or item.get("used_in_word") is False:
            continue
        raw_path = item.get("file") or item.get("path") or item.get("filename")
        if not raw_path:
            continue
        path = Path(str(raw_path))
        if not path.is_absolute():
            path = manifest_path.parent / path
        if path.exists():
            order = int(item.get("order") or item.get("sort") or index)
            step = str(item.get("flow_step") or item.get("label") or path.name)
            selected.append((order, step, path))
    selected.sort(key=lambda row: (row[0], row[1], row[2].name))
    return [path for _, _, path in selected]


def screenshot_dimensions(path: Path) -> tuple[int, int] | None:
    if Image is None:
        return None
    try:
        with Image.open(path) as img:
            return img.size
    except Exception:
        return None


def is_complete_screenshot(path: Path) -> bool:
    dims = screenshot_dimensions(path)
    if dims is None:
        return True
    width, height = dims
    if width < 900 or height < 600:
        return False
    if width / max(height, 1) > 2.25:
        return False
    return True


def select_delivery_screenshots(paths: list[Path], qa: list[str]) -> list[Path]:
    valid = [path for path in paths if is_complete_screenshot(path)]
    word_captures = [path for path in valid if path.name.lower().startswith("word_")]
    if word_captures:
        ignored = len(paths) - len(word_captures)
        qa.append(f"PASS: screenshot filter selected {len(word_captures)} full-window word_ capture(s); ignored {ignored} legacy/cropped image(s)")
        return word_captures
    if len(valid) != len(paths):
        qa.append(f"PASS: screenshot filter removed {len(paths) - len(valid)} cropped or undersized image(s)")
    return valid


def cn_number(num: int) -> str:
    digits = "零一二三四五六七八九"
    if 1 <= num <= 10:
        return "十" if num == 10 else digits[num]
    if 11 <= num <= 19:
        return "十" + digits[num - 10]
    if 20 <= num <= 99:
        tens, ones = divmod(num, 10)
        return digits[tens] + "十" + (digits[ones] if ones else "")
    return str(num)


def extract_manual_modules(business: dict[str, str]) -> list[tuple[str, str]]:
    text = business.get("functions") or ""
    modules: list[tuple[str, str]] = []
    for raw in text.splitlines():
        line = raw.strip(" ・-;；")
        if not line:
            continue
        if "：" in line:
            title, body = line.split("：", 1)
        elif ":" in line:
            title, body = line.split(":", 1)
        else:
            title, body = line[:18], line
        title = re.sub(r"^[（(]?\d+[）)]?\s*", "", title).strip()
        body = body.strip() or line
        if not title or len(title) > 28:
            title = "功能模块操作"
        modules.append((title, body))
    if modules:
        return modules[:6]
    return [
        ("数据录入与配置", "用户在该模块录入或选择待处理对象，设置必要参数，完成处理前的基础配置。"),
        ("业务处理与结果查看", "用户触发处理操作后，软件按照预设逻辑完成计算、整理、筛选、生成或校验，并在界面中展示处理结果和状态提示。"),
        ("记录管理与资料导出", "用户可查看历史记录、导出结果文件或生成报告，便于后续复核、归档和交付。"),
    ]


def parse_output_rows(manual_text: str) -> list[tuple[str, str]]:
    rows: list[tuple[str, str]] = []
    in_table = False
    for raw in manual_text.splitlines():
        line = raw.strip()
        if not line.startswith("|") or "|" not in line[1:]:
            if in_table:
                break
            continue
        cells = [clean_value(cell) for cell in line.strip("|").split("|")]
        if len(cells) < 2:
            continue
        if all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            in_table = True
            continue
        if cells[0] in {"字段", "名称", "输出项"}:
            in_table = True
            continue
        if in_table:
            rows.append((cells[0], cells[1]))
    return rows[:12]


def render_mermaid_blocks(blocks: list[str], out_dir: Path, qa: list[str]) -> list[Path]:
    if not blocks:
        return []
    out_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    for index, block in enumerate(blocks, start=1):
        path = out_dir / f"mermaid_flow_{index:02d}.png"
        try:
            draw_mermaid_flowchart(block, path)
            paths.append(path)
        except Exception as exc:
            qa.append(f"WARN: failed to render Mermaid block {index}: {exc}")
    if paths:
        qa.append(f"PASS: rendered {len(paths)} Mermaid diagram(s) to images")
    return paths


def pick_screenshots(screenshots: list[Path], keywords: list[str], used: set[Path], *, fallback: bool = True, mark_used: bool = True) -> list[Path]:
    for keyword in keywords:
        for path in screenshots:
            if path in used:
                continue
            if keyword.lower() in path.name.lower():
                if mark_used:
                    used.add(path)
                return [path]
    if fallback:
        for path in screenshots:
            if path not in used:
                if mark_used:
                    used.add(path)
                return [path]
    return []


def module_screenshot_keywords(title: str) -> list[str]:
    mapping = [
        ("启动", ["start", "window", "main"]),
        ("主界面", ["start", "window", "main", "overview"]),
        ("添加", ["sample_file_added", "add", "file"]),
        ("导入", ["sample_file_added", "input", "file"]),
        ("文件", ["sample_file_added", "file"]),
        ("转换", ["conversion_log_visible", "conversion_log", "convert", "log"]),
        ("处理", ["conversion_log_visible", "conversion_log", "process", "log"]),
        ("日志", ["conversion_log_visible", "conversion_log", "log"]),
        ("输出", ["output", "result", "report", "package"]),
        ("质检", ["report", "package", "conversion_log"]),
        ("申报", ["report", "package", "conversion_log"]),
        ("单点", ["main_evaluation_result", "evaluation_result"]),
        ("批量", ["batch_evaluation", "batch"]),
        ("曲线", ["model_curves", "curve"]),
        ("记录", ["history_records", "history"]),
        ("报告", ["generated_html_report", "report"]),
        ("工况", ["single_case_workspace", "workspace"]),
        ("录入", ["single_case_workspace", "workspace"]),
        ("评估", ["main_evaluation_result", "evaluation_result"]),
        ("移动", ["mobile_responsive", "mobile"]),
    ]
    for needle, keywords in mapping:
        if needle in title:
            return keywords
    return []


def parse_mermaid_node(token: str) -> tuple[str, str]:
    token = token.strip().strip(";")
    token = re.sub(r"^&\s*", "", token)
    match = re.match(r"([A-Za-z0-9_\u4e00-\u9fff-]+)\s*(?:\[(.*?)\]|\((.*?)\)|\{(.*?)\})?", token)
    if not match:
        label = token.strip("\"'")
        return label, label
    node_id = match.group(1)
    label = next((group for group in match.groups()[1:] if group), None) or node_id
    return node_id, label.strip("\"'")


def parse_mermaid_edges(block: str) -> tuple[str, dict[str, str], list[tuple[str, str, str]]]:
    direction = "TD"
    nodes: dict[str, str] = {}
    edges: list[tuple[str, str, str]] = []
    for raw in block.splitlines():
        line = raw.strip()
        if not line or line.startswith("%%"):
            continue
        header = re.match(r"^(flowchart|graph)\s+(TD|TB|BT|LR|RL)", line, re.IGNORECASE)
        if header:
            direction = header.group(2).upper()
            continue
        if "-->" not in line and "---" not in line:
            node_id, label = parse_mermaid_node(line)
            if node_id:
                nodes.setdefault(node_id, label)
            continue
        label = ""
        line = line.rstrip(";")
        label_match = re.search(r"-->\|([^|]+)\|", line)
        if label_match:
            label = label_match.group(1).strip()
            left, right = re.split(r"-->\|[^|]+\|", line, maxsplit=1)
        elif "-->" in line:
            left, right = line.split("-->", 1)
        else:
            left, right = line.split("---", 1)
        left_id, left_label = parse_mermaid_node(left)
        right_id, right_label = parse_mermaid_node(right)
        nodes.setdefault(left_id, left_label)
        nodes.setdefault(right_id, right_label)
        edges.append((left_id, right_id, label))
    return direction, nodes, edges


def load_diagram_font(size: int):
    if ImageFont is None:
        return None
    font_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    for name in ["msyh.ttc", "simhei.ttf", "simsun.ttc", "arial.ttf"]:
        path = font_dir / name
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text_for_box(text: str, width: int = 12) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if len(text) <= width:
        return text
    return "\n".join(text[i : i + width] for i in range(0, len(text), width))


def draw_mermaid_flowchart(block: str, path: Path) -> None:
    if Image is None or ImageDraw is None:
        raise RuntimeError("Pillow is required to render Mermaid diagrams")
    direction, nodes, edges = parse_mermaid_edges(block)
    if not nodes:
        raise ValueError("no Mermaid flowchart nodes parsed")
    ordered = list(nodes.keys())
    horizontal = direction in {"LR", "RL"}
    box_w, box_h = 190, 76
    gap_x, gap_y = 70, 50
    margin = 50
    if horizontal:
        width = margin * 2 + len(ordered) * box_w + max(0, len(ordered) - 1) * gap_x
        height = margin * 2 + box_h + 40
        positions = {node: (margin + idx * (box_w + gap_x), margin + 20) for idx, node in enumerate(ordered)}
    else:
        width = margin * 2 + box_w + 320
        height = margin * 2 + len(ordered) * box_h + max(0, len(ordered) - 1) * gap_y
        positions = {node: (margin + 160, margin + idx * (box_h + gap_y)) for idx, node in enumerate(ordered)}
    img = Image.new("RGB", (max(width, 640), max(height, 360)), "white")
    draw = ImageDraw.Draw(img)
    font = load_diagram_font(20)
    small_font = load_diagram_font(16)
    border = (44, 88, 156)
    fill = (239, 246, 255)
    line_color = (60, 70, 85)

    for src, dst, label in edges:
        if src not in positions or dst not in positions:
            continue
        x1, y1 = positions[src]
        x2, y2 = positions[dst]
        start = (x1 + box_w if horizontal else x1 + box_w // 2, y1 + box_h // 2 if horizontal else y1 + box_h)
        end = (x2 if horizontal else x2 + box_w // 2, y2 + box_h // 2 if horizontal else y2)
        draw.line([start, end], fill=line_color, width=3)
        arrow = 10
        if horizontal:
            draw.polygon([(end[0], end[1]), (end[0] - arrow, end[1] - arrow // 2), (end[0] - arrow, end[1] + arrow // 2)], fill=line_color)
        else:
            draw.polygon([(end[0], end[1]), (end[0] - arrow // 2, end[1] - arrow), (end[0] + arrow // 2, end[1] - arrow)], fill=line_color)
        if label:
            mid = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2)
            draw.text((mid[0] + 4, mid[1] - 18), label, fill=(90, 90, 90), font=small_font)

    for node, label in nodes.items():
        x, y = positions[node]
        draw.rounded_rectangle([x, y, x + box_w, y + box_h], radius=12, fill=fill, outline=border, width=3)
        wrapped = wrap_text_for_box(label)
        bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=4)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
        draw.multiline_text((x + (box_w - text_w) / 2, y + (box_h - text_h) / 2 - 2), wrapped, fill=(20, 30, 45), font=font, align="center", spacing=4)
    img.save(path)


def insert_diagram_block(doc: Document, diagrams: list[Path], label: str, qa: list[str]) -> None:
    for path in diagrams:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"图：{label}（{path.name}）")
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(9)
        try:
            doc.add_picture(str(path), width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            qa.append(f"WARN: failed to insert diagram {path}: {exc}")


def verify_manual_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    text = collect_document_text(doc)
    checks: list[str] = []
    checks.extend(text_quality_issues("manual docx text", text, min_cjk=120))
    forbidden = ["附录：原始操作手册草稿摘要", "```mermaid", "```", "草稿摘要"]
    hits = [item for item in forbidden if item in text]
    if hits:
        checks.append("FAIL: manual contains draft/code artifacts: " + "、".join(hits))
    if re.search(r"(?m)^\s*#{1,6}\s+", text):
        checks.append("FAIL: manual contains raw Markdown heading markers")
    image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    if image_count == 0:
        checks.append("FAIL: manual contains no embedded screenshots/images")
    elif image_count < 3:
        checks.append(f"WARN: manual image count is low, count={image_count}")
    return checks or ["PASS: manual has no raw draft appendix or Mermaid code blocks"]


def manual_heading_text(section_no: int, title: str, numbered_headings: bool) -> str:
    return f"{cn_number(section_no)}、{title}" if numbered_headings else title


def build_manual_docx(
    workdir: Path,
    out_dir: Path,
    meta: dict[str, str],
    business: dict[str, str],
    screenshots_dir: Path | None,
    *,
    min_screenshots: int = 7,
    numbered_headings: bool = True,
) -> tuple[GeneratedFile, list[str]]:
    qa: list[str] = []
    out_path = out_dir / f"03 {safe_filename(meta['software_name'] + meta['version'])} 操作说明书.docx"
    all_screenshots = find_screenshots(workdir, screenshots_dir)
    screenshots = select_delivery_screenshots(all_screenshots, qa)
    manual_text = read_text(workdir / "草稿" / "操作手册.md")
    diagrams = render_mermaid_blocks(extract_mermaid_blocks(manual_text), out_dir / "diagrams", qa)
    modules = extract_manual_modules(business)

    doc = Document()
    setup_normal_style(doc)
    set_header_footer(doc, f"{meta['software_name']}{meta['version']} 操作说明书")
    add_cover(doc, meta, "操  作  说  明  书", f"编制人：{meta['developer']}    版本：{meta['version']}    日期：{datetime.now().strftime('%Y年%m月%d日')}")

    section_no = 1
    add_heading_cn(doc, manual_heading_text(section_no, "软件简介", numbered_headings), 1)
    add_paragraphs(doc, compose_intro(meta, business))
    section_no += 1
    add_heading_cn(doc, manual_heading_text(section_no, "运行环境", numbered_headings), 1)
    add_paragraphs(doc, business.get("environment") or "本软件运行于可使用 Python 3.10 以上环境的个人计算机，通过本地桌面界面或命令行入口完成操作。")
    section_no += 1
    add_heading_cn(doc, manual_heading_text(section_no, "软件启动与主界面总览", numbered_headings), 1)
    add_paragraphs(doc, "用户按照软件部署说明启动程序后进入主界面。主界面通常包括功能入口、参数或数据录入区、处理结果区、记录区和导出入口；具体控件以实际截图为准。")
    used_screenshots: set[Path] = set()
    insert_screenshot_block(
        doc,
        pick_screenshots(screenshots, ["start", "window", "main", "overview", "single_case_workspace", "main_evaluation_result"], used_screenshots),
        "主界面总览",
        qa,
        required=True,
    )

    if diagrams:
        section_no += 1
        add_heading_cn(doc, manual_heading_text(section_no, "业务流程图", numbered_headings), 1)
        add_paragraphs(doc, "本节流程图由操作手册草稿中的 Mermaid 流程描述转换为图片后插入，用于展示软件主要操作路径。")
        insert_diagram_block(doc, diagrams, "业务流程图", qa)

    for title, body in modules:
        section_no += 1
        add_heading_cn(doc, manual_heading_text(section_no, title, numbered_headings), 1)
        add_paragraphs(doc, body)
        chunk = pick_screenshots(screenshots, module_screenshot_keywords(title), used_screenshots)
        insert_screenshot_block(doc, chunk, title, qa)

    remaining_screenshots = [path for path in screenshots if path not in used_screenshots]
    if remaining_screenshots:
        section_no += 1
        add_heading_cn(doc, manual_heading_text(section_no, "补充操作截图", numbered_headings), 1)
        add_paragraphs(doc, "以下截图用于补充展示软件运行过程中的文件导入、处理反馈、日志显示或输出结果等界面状态。")
        insert_screenshot_block(doc, remaining_screenshots[:12], "补充操作截图", qa)

    section_no += 1
    add_heading_cn(doc, manual_heading_text(section_no, "输出结果字段说明", numbered_headings), 1)
    add_output_table(doc, parse_output_rows(manual_text))
    section_no += 1
    add_heading_cn(doc, manual_heading_text(section_no, "计算方法和技术特点", numbered_headings), 1)
    add_paragraphs(doc, compose_technical_text(meta, business))
    section_no += 1
    add_heading_cn(doc, manual_heading_text(section_no, "常见问题与处理建议", numbered_headings), 1)
    add_faq(doc)
    section_no += 1
    add_heading_cn(doc, manual_heading_text(section_no, "使用边界与注意事项", numbered_headings), 1)
    notes = business.get("notes") or "软件输出结果应结合业务规则、原始数据和人工复核使用。涉及正式提交、合同、财务、医疗、法律、工程安全等场景时，应按适用制度或标准进行人工确认。"
    add_paragraphs(doc, notes)

    doc.save(str(out_path))
    qa.append(f"{'PASS' if len(screenshots) >= min_screenshots else 'WARN'}: manual selected screenshot count={len(screenshots)}, expected>={min_screenshots}, discovered={len(all_screenshots)}")
    qa.extend(verify_manual_docx(out_path))
    return GeneratedFile("操作说明书", out_path), qa


def compose_intro(meta: dict[str, str], business: dict[str, str]) -> str:
    industry = strip_terminal_punct(meta["industry"])
    positioning = business.get("positioning") or ""
    if positioning:
        lead = positioning
        if meta["software_name"] not in lead:
            lead = f"{meta['software_name']}（简称：{meta['short_name']}，版本：{meta['version']}）是一款面向{industry}的软件。{lead}"
    else:
        lead = f"{meta['software_name']}（简称：{meta['short_name']}，版本：{meta['version']}）是一款面向{industry}的软件。"
    return (
        f"{lead}\n\n"
        "本说明书面向实际操作者，说明软件的运行环境、启动方式、主要界面、功能操作、输出结果和使用边界。"
        "用户可依据本说明书完成基础配置、数据录入、功能执行、结果查看和资料导出。"
    )


def add_paragraphs(doc: Document, text: str) -> None:
    for raw in re.split(r"\n{1,2}", text.strip()):
        line = raw.strip()
        if not line:
            continue
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(line)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def insert_screenshot_block(doc: Document, screenshots: list[Path], label: str, qa: list[str], *, required: bool = False) -> None:
    if not screenshots:
        if not required:
            return
        qa.append(f"FAIL: missing required screenshot for {label}")
        return
    for path in screenshots:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = caption.add_run(f"图：{label}（{path.name}）")
        cap_run.font.name = "宋体"
        cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        cap_run.font.size = Pt(9)
        try:
            width = Inches(6.2)
            if Image:
                with Image.open(path) as img:
                    if img.width < 900 or img.height < 450:
                        qa.append(f"WARN: screenshot may be too small or cropped: {path.name} {img.width}x{img.height}")
                    if img.width / max(img.height, 1) > 3.0:
                        qa.append(f"WARN: screenshot has very wide aspect ratio; check completeness: {path.name} {img.width}x{img.height}")
                    if img.width < img.height:
                        width = Inches(4.0)
            doc.add_picture(str(path), width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:
            qa.append(f"WARN: failed to insert screenshot {path}: {exc}")


def add_output_table(doc: Document, rows: list[tuple[str, str]] | None = None) -> None:
    rows = rows or [
        ("输入数据", "用户在界面或文件中录入的业务参数、基础数据或待处理对象。"),
        ("处理结果", "软件按照功能逻辑计算、整理、筛选或生成的结果信息。"),
        ("状态提示", "软件在处理成功、失败、异常或待补充时给出的提示信息。"),
        ("导出文件", "软件生成的表格、报告、记录或其他可归档文件。"),
        ("操作时间", "软件记录或展示的处理时间、创建时间或最近更新时间。"),
        ("备注说明", "用户或系统对当前记录、结果或报告的补充说明。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "含义"
    for field, meaning in rows:
        row = table.add_row()
        row.cells[0].text = field
        row.cells[1].text = meaning
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "宋体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                    run.font.size = Pt(9.5)


def add_faq(doc: Document) -> None:
    items = [
        ("软件无法启动怎么办？", "检查运行环境、启动命令、端口占用和所需依赖是否满足要求。若仍无法启动，应记录错误提示并交由开发或运维人员复核。"),
        ("输入数据后结果异常怎么办？", "先核对字段单位、数据格式、必填项和取值范围，再重新运行。批量数据应重点检查表头、分隔符和空值。"),
        ("处理过程提示失败怎么办？", "根据页面或命令行提示定位失败步骤，检查输入文件是否存在、是否被其他程序占用，以及输出目录是否有写入权限。"),
        ("导出文件在哪里查看？", "导出文件通常保存在软件配置的输出目录或界面提示路径中，可按生成时间和文件名进行查找。"),
        ("历史记录没有更新怎么办？", "检查当前操作是否已执行保存，确认本地数据库或记录文件可写，并刷新记录列表。"),
        ("软件结果能否直接替代人工判断？", "不能。软件结果用于提高处理效率和一致性，正式决策仍应结合业务规则、原始资料和人工复核。"),
    ]
    for q, a in items:
        para = doc.add_paragraph()
        run = para.add_run(q)
        run.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)
        add_paragraphs(doc, a)


def collect_pending(meta: dict[str, str]) -> list[str]:
    return [f"{key}: {value}" for key, value in meta.items() if is_pending(value)]


def write_report(out_dir: Path, generated: list[GeneratedFile], meta: dict[str, str], qa_sections: dict[str, list[str]]) -> Path:
    path = out_dir / "word_generation_report.md"
    lines = [
        "# Word 三件套生成报告",
        "",
        f"- 时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 软件：{meta['software_name']}",
        f"- 版本：{meta['version']}",
        "",
        "## 生成文件",
        "",
    ]
    for item in generated:
        lines.append(f"- {item.label}：`{item.path}`")
    pending = collect_pending(meta)
    lines.extend(["", "## 待确认字段", ""])
    if pending:
        lines.extend(f"- {item}" for item in pending)
    else:
        lines.append("- 无")
    lines.extend(["", "## QA 记录", ""])
    for name, items in qa_sections.items():
        lines.append(f"### {name}")
        lines.extend(f"- {item}" for item in items)
        lines.append("")
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def build_all(args: argparse.Namespace) -> int:
    workdir = Path(args.workdir).resolve()
    reference_word_dir = Path(args.reference_word_dir).resolve() if args.reference_word_dir else None
    template_dir = Path(args.template_dir).resolve() if args.template_dir else reference_word_dir
    out_dir = Path(args.out_dir).resolve() if args.out_dir else workdir / "正式资料" / "word"
    out_dir.mkdir(parents=True, exist_ok=True)

    meta = load_metadata(workdir, args)
    business = extract_business_paragraphs(workdir)
    screenshots_dir = Path(args.screenshots_dir).resolve() if args.screenshots_dir else None
    reference_image_count = reference_manual_image_count(reference_word_dir)
    min_screenshots = reference_image_count or 7
    numbered_headings = reference_word_dir is None

    generated: list[GeneratedFile] = []
    qa_sections: dict[str, list[str]] = {}

    if reference_word_dir:
        qa_sections["参考模板"] = [
            f"PASS: reference word dir used: {reference_word_dir}",
            f"PASS: application form template source={template_dir}",
            f"PASS: manual heading style={'numbered' if numbered_headings else 'plain'}",
            f"PASS: reference manual image count={reference_image_count or 'not-detected'}",
        ]

    app, app_qa = build_application_docx(workdir, out_dir, meta, business, template_dir)
    generated.append(app)
    qa_sections["申请表"] = app_qa

    source, source_qa = build_source_docx(workdir, out_dir, meta)
    generated.append(source)
    qa_sections["源代码"] = source_qa

    manual, manual_qa = build_manual_docx(
        workdir,
        out_dir,
        meta,
        business,
        screenshots_dir,
        min_screenshots=min_screenshots,
        numbered_headings=numbered_headings,
    )
    generated.append(manual)
    qa_sections["操作说明书"] = manual_qa

    report_path = write_report(out_dir, generated, meta, qa_sections)
    has_failures = any(item.startswith("FAIL") for items in qa_sections.values() for item in items)
    print(f"{'FAIL' if has_failures else 'OK'} generated {len(generated)} DOCX files")
    print(report_path)
    for item in generated:
        print(item.path)
    return 2 if has_failures else 0


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--workdir", required=True, help="软件著作权申请资料目录")
    parser.add_argument("--template-dir", help="代理样本文档或申请表模板目录")
    parser.add_argument("--reference-word-dir", help="参考 Word 三件套目录；用于复制申请表模板、对齐说明书标题样式和截图数量")
    parser.add_argument("--out-dir", help="DOCX 输出目录，默认 <workdir>/正式资料/word")
    parser.add_argument("--screenshots-dir", help="优先使用的截图目录")
    parser.add_argument("--software-name")
    parser.add_argument("--short-name")
    parser.add_argument("--version")
    parser.add_argument("--owner")
    parser.add_argument("--developer")
    parser.add_argument("--completion-date")
    parser.add_argument("--first-publish")
    parser.add_argument("--rights-acquire")
    parser.add_argument("--development-mode")
    parser.add_argument("--source-lines")
    parser.add_argument("--industry")
    parser.add_argument("--classification-code")
    parser.add_argument("--owner-category")
    parser.add_argument("--certificate-type")
    parser.add_argument("--certificate-no")
    parser.add_argument("--nationality")
    parser.add_argument("--province-city")
    parser.add_argument("--contact-name")
    parser.add_argument("--contact-phone")
    parser.add_argument("--contact-mobile")
    parser.add_argument("--contact-email")
    parser.add_argument("--contact-address")
    parser.add_argument("--contact-postcode")
    return parser.parse_args()


if __name__ == "__main__":
    raise SystemExit(build_all(parse_args()))
